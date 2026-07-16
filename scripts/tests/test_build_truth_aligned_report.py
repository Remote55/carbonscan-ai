"""Tests for the conservative NSC report builder."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path

import pytest
from docx import Document
from scripts.build_truth_aligned_report import (
    ANCHORS,
    FORBIDDEN_CLAIMS,
    TABLE_ROW_ANCHORS,
    build_report,
)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _make_source(path: Path) -> None:
    document = Document()
    first = document.add_paragraph()
    first_run = first.add_run(f"{ANCHORS[0]} fixture legacy claim")
    first_run.bold = True
    for anchor in ANCHORS[1:]:
        document.add_paragraph(f"{anchor} fixture legacy claim")

    # Regression: this caption must not be mistaken for the narrower
    # "Allometric Carbon Calc ระบบดำเนินการ" body anchor.
    document.add_paragraph("รูปที่ 12 Allometric Carbon Calculation (คำนวณคาร์บอน)")

    table = document.add_table(rows=1, cols=4)
    table.cell(0, 0).text = "kept"
    for anchor in TABLE_ROW_ANCHORS:
        row = table.add_row()
        row.cells[0].text = anchor
        for cell in row.cells[1:]:
            cell.text = "legacy"
    document.save(path)


def _make_manifest(path: Path) -> None:
    path.write_text(
        json.dumps(
            {
                "baseline": {"backend": "tlsep", "status": "Implemented"},
                "candidate": {
                    "display_name": "PointNet++",
                    "status": "Experimental",
                },
                "validation": {
                    "synthetic_held_out": {
                        "pointnet_mean_iou": 0.977625,
                        "tlsep_mean_iou": 0.7692083333,
                    },
                    "wan_held_out": {
                        "wood_iou": 0.418,
                        "leaf_iou": 0.808,
                        "mean_iou": 0.613,
                        "accuracy": 0.831,
                    },
                    "demol_65": {
                        "dbh_mae_cm": 1.1673846154,
                        "height_mae_m": 0.5446153846,
                        "volume_mape_pct": 18.7650916186,
                    },
                },
            }
        ),
        encoding="utf-8",
    )


def test_builder_never_changes_source_and_preserves_structure(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    manifest = tmp_path / "manifest.json"
    _make_source(source)
    _make_manifest(manifest)
    source_before = _sha256(source)

    audit = build_report(source, output, manifest)

    assert _sha256(source) == source_before
    assert output.is_file()
    assert audit.source_sha256 == source_before
    assert audit.source_unchanged is True
    assert audit.anchors_replaced == len(ANCHORS) + len(TABLE_ROW_ANCHORS)
    assert audit.paragraph_anchors_replaced == len(ANCHORS)
    assert audit.table_row_anchors_replaced == len(TABLE_ROW_ANCHORS)
    assert audit.truth_contract_passed is True
    assert audit.tables_before == audit.tables_after == 1
    assert audit.inline_shapes_before == audit.inline_shapes_after == 0
    assert audit.sections_before == audit.sections_after == 1
    assert audit.page_geometry_preserved is True

    built = Document(output)
    assert built.paragraphs[0].runs[0].bold is True
    all_text = "\n".join(paragraph.text for paragraph in built.paragraphs)
    assert "[Implemented]" in all_text
    assert "[Experimental]" in all_text
    assert "[Planned]" in all_text
    assert "Wood IoU 0.418" in all_text
    assert "DBH MAE 1.1673846154 cm" in all_text
    assert "ประยุกต์ใช้ ResNet50" not in all_text
    assert "ความแม่นยำได้สูงถึง 0.978" not in all_text

    table_text = "\n".join(
        cell.text for table in built.tables for row in table.rows for cell in row.cells
    )
    assert "[Implemented default] tlsep; [Experimental] PointNet++" in table_text
    assert "0.418" in table_text
    assert "0.808" in table_text
    assert "0.613" in table_text
    assert "[Stub] Species classifier" in table_text


def test_builder_aborts_when_anchor_is_missing(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    manifest = tmp_path / "manifest.json"
    _make_source(source)
    _make_manifest(manifest)
    document = Document(source)
    document.paragraphs[1].text = "anchor removed"
    document.save(source)

    with pytest.raises(ValueError, match="PointNet"):
        build_report(source, output, manifest)

    assert not output.exists()


def test_builder_rejects_source_as_output(tmp_path: Path):
    source = tmp_path / "source.docx"
    manifest = tmp_path / "manifest.json"
    _make_source(source)
    _make_manifest(manifest)

    with pytest.raises(ValueError, match="different paths"):
        build_report(source, source, manifest)


def test_builder_aborts_when_forbidden_claim_survives(tmp_path: Path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    manifest = tmp_path / "manifest.json"
    _make_source(source)
    _make_manifest(manifest)
    document = Document(source)
    document.add_paragraph(FORBIDDEN_CLAIMS[7])
    document.save(source)

    with pytest.raises(ValueError, match="Forbidden legacy claim"):
        build_report(source, output, manifest)

    assert not output.exists()
