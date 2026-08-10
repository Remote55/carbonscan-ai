"""Build a truth-aligned NSC report without modifying the source DOCX."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import shutil
import tempfile
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.table import _Row
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class ReportAudit:
    """Machine-readable proof that the source and document structure survived."""

    source_path: str
    output_path: str
    source_sha256: str
    output_sha256: str
    source_unchanged: bool
    anchors_replaced: int
    paragraph_anchors_replaced: int
    table_row_anchors_replaced: int
    truth_contract_passed: bool
    paragraphs_before: int
    paragraphs_after: int
    tables_before: int
    tables_after: int
    inline_shapes_before: int
    inline_shapes_after: int
    media_files_before: int
    media_files_after: int
    sections_before: int
    sections_after: int
    page_geometry_preserved: bool


@dataclass(frozen=True)
class _Structure:
    paragraphs: int
    tables: int
    inline_shapes: int
    media_files: int
    sections: int
    page_geometry: tuple[tuple[Any, ...], ...]


ANCHORS = (
    "Tree Segmentation เมื่อได้รับแบบจำลอง CHM แล้ว",
    "Wood-Leaf Separation ในขั้นตอนดังกล่าว ระบบจะใช้ Deep Learning เช่น PointNet++",
    "QSM (Cylinder Fitting)",
    "Allometric Carbon Calc ระบบดำเนินการ",
    "ผลบนชุดทดสอบ PointNet++",
    "6.5.2.1 การทดสอบ Wood-Leaf บนไม้จริง",
    "ระบบอัปโหลดและการประมวลผลแบบอะซิงโครนัส",
    "ระบบแผนที่ภูมิสารสนเทศ (GIS Map)",
    "ตลาดกลางคาร์บอนเครดิต (Marketplace)",
    "การออกใบรับรองและการซื้อขาย",
    "[7] TGO",
    "[18] Demol",
    "[20] Wan",
    "รูปที่ 2 GIS Map แสดงผลแปลงปลูกป่า",
    "Ground Classification ระบบจะดำเนินการจำแนกจุด",
    "Canopy Height Model ระบบจะสร้างแผนที่ความสูง",
    "Species Classification ระบบประยุกต์ใช้ ResNet50",
    "บทบาท: ใช้พัฒนาระบบ Backend",
    "ระบบทำธุรกรรม: การส่งคำสั่งซื้อคาร์บอนเครดิต",
    "โครงสร้าง Backend และ API:",
    "โครงสร้าง Machine Learning Pipeline:",
    "ตารางการทำงานและธุรกรรม:",
    "ระบบใช้สถาปัตยกรรมแบบ Asynchronous",
    "การประเมินประสิทธิภาพในการแยกส่วนประกอบของต้นไม้",
    "โมเดล PointNet++ สามารถทำค่าความแม่นยำได้สูงถึง 0.978",
    "ความคลาดเคลื่อนของปัญญาประดิษฐ์กับลักษณะป่าไม้ไทย:",
    "โครงงาน TreeQ Carbon Platform: ระบบวิเคราะห์และคำนวณการกักเก็บคาร์บอนจากต้นไม้ เป็นระบบที่พัฒนาขึ้น",
    "ระบบสามารถรับข้อมูลของต้นไม้ เช่น ชนิด ขนาด แล้วนำมาวิเคราะห์",
    "ดังนั้น เพื่อแก้ไขปัญหาดังกล่าว คณะผู้จัดทำจึงได้ริเริ่มแนวคิด",
    "3.2 เพิ่มความแม่นยำในการวิเคราะห์ด้วยเทคโนโลยี AI",
    "โครงงาน TreeQ Carbon Platform เป็นระบบที่พัฒนาขึ้นเพื่อช่วยวิเคราะห์และคำนวณ",
    "ระบบถูกออกแบบให้ช่วยลดความซับซ้อนของการคำนวณแบบดั้งเดิม",
    "Height Normalization ภายหลังจากการจำแนกพื้นดิน",
    "รูปที่ 11 Species Classification (ระบุชนิดพันธุ์)",
    "โดยค่าคงที่และสมการทั้งหมดอ้างอิงจาก Chave et al. (2014)",
    "ปัญญาประดิษฐ์ไม่ได้ประเมินเพียงจุดใดจุดหนึ่งแยกส่วนกัน",
    "บทบาท: ใช้พัฒนาระบบปัญญาประดิษฐ์ และ Machine Learning",
    "Platform: Supabase (PostgreSQL 16)",
    "วัตถุประสงค์: ใช้เป็นระบบฐานข้อมูลกลางของโครงงาน",
    "Platform: RunPod Serverless",
    "วัตถุประสงค์: ใช้สำหรับประมวลผล AI และ Machine Learning ที่ต้องการทรัพยากร",
    "ข้อมูลที่นำเข้ามาจะถูกส่งไปยังระบบคลาวด์ที่มีหน่วยประมวลผล",
    "การแสดงผลเชิงภาพ: ผู้ใช้สามารถดูโมเดลต้นไม้ของตนเอง",
    "การจัดการบัญชีผู้ใช้และพื้นที่แปลงปลูก:",
    "ระบบความปลอดภัยและการเข้าถึงข้อมูล:",
    "ตารางผู้ใช้และแปลงป่า:",
    "ตารางข้อมูลแกนหลัก:",
    "ตารางข้อมูลอ้างอิง: แหล่งรวมข้อมูลตัวแปรความหนาแน่นเนื้อไม้",
    "การทำงานของ REST API Endpoints:",
    "ผลการตรวจสอบความแม่นยำ: เมื่อเปรียบเทียบค่าที่ระบบวัดได้",
    "ผลการทดสอบพบว่าแบบจำลองที่ฝึกด้วยข้อมูลสังเคราะห์ล้วน",
    "รูปที่ 16 หน้าจอแสดงสถานะกำลังประมวลผลความคืบหน้า",
    "ในส่วนของการตรวจสอบความถูกต้องกับชุดข้อมูลไม้จริง",
    "เป้าหมายที่ 13: ช่วยเตรียมการรับมือการเปลี่ยนแปลงสภาพภูมิอากาศ",
    "เป้าหมายที่ 15: ปกป้อง ฟื้นฟู และสนับสนุนการใช้ระบบนิเวศ",
    "ข้อจำกัดเรื่องพันธุ์ไม้และสมการคาร์บอน:",
    "ข้อพิจารณาเชิงชีววิทยาของไผ่:",
    "ระยะเวลาในการประมวลผล:",
    "ชีวมวลเหนือพื้นดิน (AGB): AGB = 0.0673",
    "คาร์บอนสะสมรวม (kg C): คาร์บอน = (AGB × 1.24) × 0.47",
    "ระบบใช้ข้อมูลทั้งหมด 2 ชุดตามบทบาทการทำงาน",
    "6.5.2 ชุดข้อมูลฝึกแบบจำลอง Wood-Leaf Segmentation (PointNet++)",
    "รายละเอียดชุดข้อมูล: แบบจำลองแยกส่วนลำต้นและใบ จะถูกฝึกบนชุดข้อมูล Point Cloud สังเคราะห์",
)


TABLE_ROW_ANCHORS = (
    "PCA heuristic (zero-shot)",
    "PointNet++ ฝึกด้วยข้อมูลสังเคราะห์ → ทดสอบไม้จริง (zero-shot)",
    "PointNet++ ฝึกด้วยไม้จริงโดยตรง + augment (ผลดีที่สุด)",
    "(อ้างอิง) PointNet++ บนชุดทดสอบสังเคราะห์",
    "1",
    "2",
    "3",
    "4",
    "5",
    "6",
    "7",
    "11",
    "17",
    "20",
    "21",
)


FORBIDDEN_CLAIMS = (
    "โดยใช้เทคนิค CSF (Cloth Simulation Filter)",
    "โดยประยุกต์ใช้วิธี Pit-free CHM",
    "Species Classification ระบบประยุกต์ใช้ ResNet50",
    "จัดการ REST API และ WebSocket เพื่อให้ระบบสื่อสารข้อมูลได้แบบ Real-time",
    "การส่งคำสั่งซื้อคาร์บอนเครดิตจะรับข้อมูลรูปแบบ JSON",
    "มีการเปิดช่องทาง WebSocket เพื่อส่งข้อมูลความคืบหน้า",
    "รายงานความคืบหน้าแบบเรียลไทม์ผ่าน WebSocket",
    "มีระบบแจ้งเตือนผ่านอีเมล",
    "มีประสิทธิภาพดีกว่าอย่างเห็นได้ชัด",
    "ยืนยันว่าแกนการวัดขนาดต้นไม้ของระบบมีความแม่นยำเทียบเท่าข้อมูลจริง",
    "เก็บข้อมูลภาคสนามของไม้ไทยเพิ่มเติม",
    "ความคลาดเคลื่อนจากการวัดหน้างานจริงประมาณร้อยละ 10 ถึง 20",
    "5 ถึง 30 นาทีต่อไฟล์",
    "CSF (Cloth Simulation Filter)",
    "Pit-free Canopy Height Model (CHM)",
    "TreeQSM (Least Squares Cylinder Fitting)",
    "ResNet50 (Transfer Learning)",
    "PointNet++ ฝึกด้วยไม้จริงโดยตรง + augment (ผลดีที่สุด)",
    "สมการแอลโลเมตริก (Allometric Equations) ตามมาตรฐาน TGO",
)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _length(value: Any) -> int | None:
    return None if value is None else int(value)


def _page_geometry(document: DocumentType) -> tuple[tuple[Any, ...], ...]:
    return tuple(
        (
            _length(section.page_width),
            _length(section.page_height),
            _length(section.top_margin),
            _length(section.bottom_margin),
            _length(section.left_margin),
            _length(section.right_margin),
            _length(section.header_distance),
            _length(section.footer_distance),
            _length(section.gutter),
            str(section.start_type),
        )
        for section in document.sections
    )


def _media_count(path: Path) -> int:
    with zipfile.ZipFile(path) as archive:
        return sum(name.startswith("word/media/") for name in archive.namelist())


def _structure(document: DocumentType, path: Path) -> _Structure:
    return _Structure(
        paragraphs=len(document.paragraphs),
        tables=len(document.tables),
        inline_shapes=len(document.inline_shapes),
        media_files=_media_count(path),
        sections=len(document.sections),
        page_geometry=_page_geometry(document),
    )


def _matches(document: DocumentType, anchor: str) -> list[Paragraph]:
    return [paragraph for paragraph in document.paragraphs if anchor in paragraph.text]


def _validate_anchor(document: DocumentType, anchor: str) -> Paragraph:
    matches = _matches(document, anchor)
    if len(matches) != 1:
        raise ValueError(
            f"Anchor must occur in exactly one paragraph: {anchor!r}; found {len(matches)}"
        )
    paragraph = matches[0]
    xml = paragraph._p.xml
    if "<w:drawing" in xml or "<w:object" in xml or "<w:pict" in xml:
        raise ValueError(
            f"Refusing to replace anchor paragraph containing media: {anchor!r}"
        )
    if not paragraph.runs:
        raise ValueError(f"Anchor paragraph has no replaceable run: {anchor!r}")
    return paragraph


def _replace_paragraph(paragraph: Paragraph, replacement: str) -> None:
    first_run = paragraph.runs[0]
    first_element = first_run._r
    first_run.text = replacement
    for run_element in list(paragraph._p.xpath(".//w:r")):
        if run_element is first_element:
            continue
        parent = run_element.getparent()
        if parent is not None:
            parent.remove(run_element)
    if paragraph.text != replacement:
        raise ValueError(
            "Paragraph replacement left unexpected field or hyperlink text"
        )


def replace_anchor(document: DocumentType, anchor: str, replacement: str) -> None:
    """Replace exactly one paragraph while preserving its style and first-run format."""
    paragraph = _validate_anchor(document, anchor)
    _replace_paragraph(paragraph, replacement)


def _matching_rows(document: DocumentType, anchor: str) -> list[_Row]:
    return [
        row
        for table in document.tables
        for row in table.rows
        if row.cells and row.cells[0].text == anchor
    ]


def _validate_table_row(document: DocumentType, anchor: str) -> _Row:
    matches = _matching_rows(document, anchor)
    if len(matches) != 1:
        raise ValueError(
            f"Table row anchor must occur exactly once: {anchor!r}; found {len(matches)}"
        )
    row = matches[0]
    for cell in row.cells:
        if len(cell.paragraphs) != 1:
            raise ValueError(
                f"Table row anchor contains a multi-paragraph cell: {anchor!r}"
            )
        paragraph = cell.paragraphs[0]
        xml = paragraph._p.xml
        if "<w:drawing" in xml or "<w:object" in xml or "<w:pict" in xml:
            raise ValueError(
                f"Refusing to replace table row containing media: {anchor!r}"
            )
        if not paragraph.runs:
            raise ValueError(f"Table row cell has no replaceable run: {anchor!r}")
    return row


def _replace_table_row(
    row: _Row,
    replacements: dict[int, str],
) -> None:
    for index, replacement in replacements.items():
        if index >= len(row.cells):
            raise ValueError(
                f"Table row has {len(row.cells)} cells; cannot replace cell {index}"
            )
        _replace_paragraph(row.cells[index].paragraphs[0], replacement)


def _require(mapping: dict[str, Any], key: str, label: str) -> Any:
    if key not in mapping:
        raise ValueError(f"Manifest {label} missing required key: {key}")
    return mapping[key]


def _replacements(manifest: dict[str, Any]) -> dict[str, str]:
    validation = _require(manifest, "validation", "root")
    wan = _require(validation, "wan_held_out", "validation")
    synthetic = _require(validation, "synthetic_held_out", "validation")
    demol = _require(validation, "demol_65", "validation")
    baseline = _require(manifest, "baseline", "root")
    candidate = _require(manifest, "candidate", "root")

    wood_iou = _require(wan, "wood_iou", "wan_held_out")
    leaf_iou = _require(wan, "leaf_iou", "wan_held_out")
    mean_iou = _require(wan, "mean_iou", "wan_held_out")
    accuracy = _require(wan, "accuracy", "wan_held_out")
    synthetic_iou = _require(synthetic, "pointnet_mean_iou", "synthetic_held_out")
    synthetic_tlsep_iou = _require(synthetic, "tlsep_mean_iou", "synthetic_held_out")
    dbh_mae = _require(demol, "dbh_mae_cm", "demol_65")
    height_mae = _require(demol, "height_mae_m", "demol_65")
    volume_mape = _require(demol, "volume_mape_pct", "demol_65")
    baseline_name = _require(baseline, "backend", "baseline")
    candidate_name = _require(candidate, "display_name", "candidate")

    return {
        ANCHORS[0]: (
            "[Implemented] Tree Segmentation เมื่อได้รับแบบจำลอง CHM แบบ max-Z ร่วมกับ "
            "morphological filling แล้ว ระบบหา local maxima เป็นตำแหน่งยอดไม้และใช้ watershed "
            "บน negative CHM เพื่อแบ่งขอบเขตต้นไม้รายต้น จากนั้น map label กลับไปยัง point cloud "
            "โดยจุดที่ต่ำกว่าเกณฑ์จะยังเป็น unassigned"
        ),
        ANCHORS[1]: (
            f"[Implemented] Wood-Leaf Separation ในขั้นตอนดังกล่าวใช้ {baseline_name} ซึ่งเป็น "
            f"local-PCA geometric baseline เป็นค่าเริ่มต้นและไม่ต้องใช้ checkpoint ส่วน "
            f"[Experimental] {candidate_name} มี backend และผลทดลอง แต่จะไม่ถูกโปรโมตจนกว่า "
            "independent real-data test จะแสดงว่า Wood IoU ดีขึ้นและ DBH/height/volume ไม่แย่ลง "
            "พร้อม checkpoint และ training provenance ที่ตรวจสอบได้"
        ),
        ANCHORS[2]: (
            "[Implemented] QSM (Cylinder Fitting) ในโค้ดปัจจุบันเป็น QSM-derived geometry: "
            "วัด DBH ด้วย RANSAC circle ที่ระดับ 1.3 เมตร วัดความสูงด้วยค่า max-Z และคำนวณ "
            "stem volume ด้วย taper equation (form factor 0.50) โดย branch volume ยังคงเป็น 0 "
            "จึงไม่ใช่ full TreeQSM และยังไม่มี branch-axis cylinder model"
        ),
        ANCHORS[3]: (
            f"[Implemented] Allometric Carbon Calc ใช้ species_db.csv เมื่อมี coefficients ครบ "
            f"และใช้ Chave 2014 fallback เมื่อไม่ทราบชนิดไม้ ผล Demol isolated-tree 65 ต้นให้ "
            f"DBH MAE {dbh_mae} cm, Height MAE {height_mae} m และ Volume MAPE {volume_mape}% "
            "ภายใต้การจำกัด 20,000 points, min-Z normalization และ tlsep ผลนี้ตรวจ geometry "
            "เท่านั้น ไม่ได้ตรวจ 8 ขั้นครบ, biomass, carbon หรือ carbon credit และ coefficients "
            "ยังต้องเทียบกับ TGO 2017 ต้นฉบับ"
        ),
        ANCHORS[4]: (
            f"[Experimental] ผลบนชุดทดสอบ PointNet++ มี recorded Mean IoU {synthetic_iou} "
            "บน held-out synthetic benchmark เท่านั้น ตัวเลขนี้แสดงว่าโมเดลเรียน synthetic "
            "distribution ได้ แต่ไม่ใช่ความแม่นบนต้นไม้จริงและใช้เป็นหลักฐานโปรโมต production "
            "default ไม่ได้"
        ),
        ANCHORS[5]: (
            f"[Experimental] 6.5.2.1 การทดสอบ Wood-Leaf บนไม้จริงจาก Wan et al. (2021) "
            f"บันทึก Wood IoU {wood_iou}, Leaf IoU {leaf_iou}, Mean IoU {mean_iou} และ "
            f"accuracy {accuracy} บน spatial held-out loader อย่างไรก็ตาม loader เดียวกันถูกใช้ "
            "เลือก best epoch และ checkpoint/tree-ID provenance สำหรับ independent final test "
            "ยังไม่ครบ จึงยังไม่โปรโมต PointNet++"
        ),
        ANCHORS[6]: (
            "[Implemented] การอัปโหลดและประมวลผลทำงานแบบ synchronous ผ่าน POST /upload/analyze "
            "คืนผลลัพธ์เต็มในคำตอบเดียว งานจริงสั้นพอ: pipeline วัดแปลง 16 ต้น 447,089 จุดใน 10 วินาที "
            "และบริการจำกัดการวิเคราะห์ไว้ที่ 200,000 จุด [Planned] คิวแบบ asynchronous เคยมีโค้ดอยู่ "
            "แต่ไม่มีผู้เรียกและไม่มี deployment ใดสตาร์ท worker จึงถูกถอดออก หากต้องการภายหลัง "
            "ต้องแก้เรื่อง object storage handoff ก่อน เพราะไฟล์อัปโหลดเคยลงดิสก์ใน container"
        ),
        ANCHORS[7]: (
            "[Planned] ระบบแผนที่ภูมิสารสนเทศ (GIS Map) และ spatial query ยังไม่ผ่าน reviewed "
            "implementation ใน prototype ปัจจุบัน หน้าเว็บมี 3D viewer แต่ห้ามนับว่า GIS พร้อมใช้"
        ),
        ANCHORS[8]: (
            "[Planned] ตลาดกลางคาร์บอนเครดิต (Marketplace), Payment Gateway และ transaction "
            "workflow ยังไม่ได้ implement ระบบปัจจุบันรายงานค่าประมาณ carbon stock/CO2e "
            "เพื่อการวิเคราะห์เท่านั้น ไม่ใช่สินทรัพย์เครดิตที่พร้อมซื้อขาย"
        ),
        ANCHORS[9]: (
            "[Planned] การออกใบรับรองและการซื้อขายยังอยู่นอกขอบเขต prototype ไม่มี MRV "
            "registry verification, project certification, credit issuance หรือ retirement "
            "ดังนั้น PDF/result ของระบบห้ามเรียกว่า certified carbon credit"
        ),
        ANCHORS[10]: (
            "[7] [Planned verification] TGO. (2017). Forestry Sector Greenhouse Gas "
            "Calculation Guideline. Thailand Greenhouse Gas Management Organization "
            "(Public Organization). https://www.tgo.or.th/ หมายเหตุ: coefficients ใน "
            "species_db.csv ยังต้องตรวจเทียบกับเอกสารต้นฉบับก่อนเคลมความสอดคล้อง"
        ),
        ANCHORS[11]: (
            "[18] [Implemented geometry evidence] Demol, M., Verbeeck, H., Gielen, B., et al. "
            "(2021). Estimating forest above-ground biomass with terrestrial laser scanning: "
            "current status and future directions. Trees, 35, 671-685. "
            "https://doi.org/10.1007/s00468-020-02067-7 Dataset: "
            "https://doi.org/10.5281/zenodo.4557401"
        ),
        ANCHORS[12]: (
            "[20] [Experimental wood-leaf evidence] Wan, P., Zhang, W., & Jin, S. (2021). "
            "Plot-level wood-leaf separation for TLS point clouds [Data set]. Dryad. "
            "https://doi.org/10.5061/dryad.rfj6q5799"
        ),
        ANCHORS[13]: (
            "[Planned] รูปที่ 2 ภาพแนวคิด GIS Map สำหรับแสดงผลแปลงปลูกป่า "
            "(ยังไม่ใช่หน้าจอ GIS ที่ผ่านการ implement และตรวจสอบใน prototype ปัจจุบัน)"
        ),
        ANCHORS[14]: (
            "[Implemented] Ground Classification จำแนก ground/non-ground ด้วย percentile-grid "
            "ground heuristic ตามโค้ดปัจจุบัน ไม่ได้ใช้ Cloth Simulation Filter (CSF) "
            "จึงไม่ควรอ้างคุณสมบัติหรือความแม่นยำของ CSF กับผลของระบบนี้"
        ),
        ANCHORS[15]: (
            "[Implemented] Canopy Height Model สร้างกริดจากค่า maximum Z และเติม/ทำให้ช่องว่างเรียบ "
            "ด้วย morphological operations ตามโค้ดปัจจุบัน วิธีนี้ไม่ใช่ full multi-threshold "
            "Pit-free CHM ของ Khosravipour et al. (2014)"
        ),
        ANCHORS[16]: (
            "[Stub] Species Classification มีเฉพาะ interface placeholder; load_model() และ "
            "classify() ยัง raise NotImplementedError และ pipeline ใช้ชนิดไม้ที่ caller ส่งเข้ามา "
            "หรือใช้สมการ Chave fallback เมื่อไม่มีชนิดไม้ [Planned] ResNet/TFLite classifier "
            "ยังต้องฝึก ประเมิน และเชื่อมเข้าระบบจริง"
        ),
        ANCHORS[17]: (
            "[Implemented] บทบาทของ Backend ปัจจุบันคือ FastAPI REST สำหรับการวิเคราะห์แบบ synchronous "
            "และ async job พร้อม worker แยก process โดยหน้าเว็บติดตามสถานะด้วย GET polling "
            "[Planned] WebSocket real-time progress ยังไม่ได้ implement"
        ),
        ANCHORS[18]: (
            "[Planned] ระบบทำธุรกรรม การส่งคำสั่งซื้อคาร์บอนเครดิต ประวัติซื้อขาย และใบรับรองอิเล็กทรอนิกส์ "
            "ยังไม่ได้ implement ผลของระบบปัจจุบันเป็นค่าประมาณ carbon stock/CO2e "
            "ไม่ใช่เครดิตที่ผ่านการรับรองหรือพร้อมซื้อขาย"
        ),
        ANCHORS[19]: (
            "[Implemented] โครงสร้าง Backend และ API ใช้ FastAPI REST มี endpoint วิเคราะห์แบบ synchronous "
            "และ async job, persistent job state และ worker แยก process โดยติดตามผลผ่าน polling "
            "[Planned] Spatial Queries, carbon-market transaction API และ WebSocket progress "
            "ยังไม่ได้ implement ในเส้นทางที่ตรวจสอบแล้ว"
        ),
        ANCHORS[20]: (
            f"โครงสร้าง Machine Learning Pipeline มี 8 ขั้นตอน: ขั้น 1-6 และ 8 [Implemented], "
            f"ขั้น 7 Species Classification [Stub]; Wood-Leaf ใช้ {baseline_name} เป็นค่าเริ่มต้น "
            f"ส่วน {candidate_name} [Experimental] ยังไม่ผ่าน evidence gate และไม่ใช่ production default"
        ),
        ANCHORS[21]: (
            "[Planned] ตาราง jobs ถูกถอดออกพร้อมคิว asynchronous ที่ไม่มีผู้เรียก "
            "ตารางประวัติการซื้อขาย การผูกผู้ซื้อกับต้นไม้ และการอ้างอิงใบรับรอง "
            "ยังไม่ได้ implement"
        ),
        ANCHORS[22]: (
            "[Implemented] การวิเคราะห์เป็น synchronous: POST /upload/analyze คืนผลลัพธ์เต็มในคำตอบเดียว "
            "[Planned] คิว asynchronous, WebSocket, ETA, email notification และ automatic retry "
            "ยังไม่ได้ implement คิวเดิมถูกถอดออกเพราะตอบ 202 queued ให้งานที่รันไม่ได้"
        ),
        ANCHORS[23]: (
            "[Experimental] การเปรียบเทียบ PCA heuristic กับ PointNet++ ในส่วนนี้เป็น synthetic held-out "
            "benchmark เท่านั้น ใช้ตรวจว่าระบบเรียนรู้ synthetic distribution ได้ แต่ไม่ใช่หลักฐาน "
            "ความแม่นยำบนต้นไม้จริงหรือหลักฐานสำหรับเปลี่ยน production default"
        ),
        ANCHORS[24]: (
            f"[Experimental] บน synthetic held-out benchmark ค่า PointNet++ Mean IoU คือ "
            f"{synthetic_iou} เทียบกับ tlsep/PCA {synthetic_tlsep_iou} (ส่วนต่าง "
            f"{synthetic_iou - synthetic_tlsep_iou:.10f}) ผลนี้ห้ามสรุปว่า AI เหนือกว่า rule-based "
            "บนต้นไม้จริง และไม่ผ่านเกณฑ์ promote ค่าเริ่มต้น"
        ),
        ANCHORS[25]: (
            "[Known limitation] PointNet++ candidate ใช้ข้อมูล open dataset จากต่างประเทศและยังไม่มี "
            "independent field validation สำหรับป่าไม้ไทย จึงยังสรุปการ generalize สู่ชนิดไม้และสภาพ "
            "ป่าเขตร้อนของไทยไม่ได้ แนวทางปัจจุบันคือเพิ่มและตรวจสอบ open datasets ตามคำแนะนำอาจารย์ "
            "พร้อมเก็บ dataset/checkpoint/tree-ID provenance ให้ครบก่อนประเมินใหม่"
        ),
        ANCHORS[26]: (
            f"[Implemented prototype] TreeQ Carbon Platform วิเคราะห์ 3D point cloud เพื่อแยกไม้/ใบ "
            f"วัด DBH และความสูง แล้วคำนวณ biomass, carbon stock และ CO2e โดยใช้ {baseline_name} "
            f"เป็นค่าเริ่มต้น ผล geometry บน Demol 65 ต้นมี DBH MAE {dbh_mae} cm; "
            f"[Experimental] {candidate_name} บน Wan มี Wood IoU {wood_iou}, Leaf IoU {leaf_iou} "
            f"และ Mean IoU {mean_iou} แต่ยังไม่ผ่าน independent promotion gate "
            "[Stub] การจำแนกชนิดไม้ และผลลัพธ์ไม่ใช่ certified carbon credit"
        ),
        ANCHORS[27]: (
            "[Implemented prototype] ระบบรับไฟล์ point cloud .las/.laz/.ply ผ่านเส้นทางวิเคราะห์ "
            "สร้างผล JSON และ segmented PLY พร้อมแสดงผลใน 3D viewer และเก็บสถานะ async jobs "
            "ชนิดไม้เป็นค่าที่ caller ระบุหรือไม่ระบุได้; classifier อัตโนมัติยังเป็น [Stub]"
        ),
        ANCHORS[28]: (
            "โครงงานจึงพัฒนาเส้นทางต้นแบบจาก 3D point cloud สู่ค่าประมาณ carbon stock/CO2e "
            "เพื่อช่วยลดงานคำนวณซ้ำและทำให้สมมติฐานตรวจสอบได้ เป้าหมายด้านภาพถ่าย GIS "
            "ตลาดคาร์บอน และการรับรองเครดิตยังเป็น [Experimental/Planned] ตามสถานะที่ระบุในเล่ม"
        ),
        ANCHORS[29]: (
            "3.2 เป้าหมาย: ประเมินว่า AI สามารถเพิ่มความแม่นยำได้หรือไม่ด้วยชุดทดสอบอิสระ "
            "ปัจจุบัน PointNet++ ยังเป็น [Experimental] และยังไม่มีหลักฐาน downstream "
            "non-regression เพียงพอสำหรับสรุปว่าแม่นกว่า baseline บนข้อมูลจริง"
        ),
        ANCHORS[30]: (
            "โครงงาน TreeQ Carbon Platform เป็น prototype วิเคราะห์ 3D point cloud ผ่าน pipeline 8 ขั้น "
            "โดยขั้น 1-6 และ 8 [Implemented], ขั้น 7 Species Classification [Stub] "
            "ผลลัพธ์เป็นค่าประมาณ biomass, carbon stock และ CO2e ที่แสดงผ่านเว็บไซต์"
        ),
        ANCHORS[31]: (
            "ระบบลดความซับซ้อนด้วยเส้นทาง demo ที่รันซ้ำได้และมี provenance/hash "
            "แต่ยังไม่ผ่าน field validation ครบ 8 ขั้นหรือการรับรองคาร์บอน "
            "การจัดเก็บและเรียกดูในปัจจุบันครอบคลุม async job/result ของเส้นทางที่ตรวจสอบแล้ว"
        ),
        ANCHORS[32]: (
            "[Implemented] Height Normalization ประมาณระดับพื้นดินของแต่ละจุดด้วย K-nearest-neighbor "
            "inverse-distance weighting จาก ground points แล้วลบค่าดังกล่าวออกจาก Z "
            "เป็น terrain normalization ตามโค้ดปัจจุบัน โดยยังไม่มี accuracy benchmark แยกสำหรับขั้นนี้"
        ),
        ANCHORS[33]: (
            "[Stub illustration] รูปที่ 11 แนวคิด Species Classification (ยังไม่มี trained classifier "
            "ที่เชื่อมกับ pipeline ปัจจุบัน)"
        ),
        ANCHORS[34]: (
            "โค้ดใช้ species_db.csv เป็น source of truth สำหรับ coefficients รายชนิดและใช้ Chave et al. "
            "(2014) เป็น pantropical fallback พร้อมค่ารากและ carbon fraction ที่บันทึกไว้ "
            "[Planned verification] ยังต้องเทียบ coefficients ทุกค่ากับเอกสาร TGO 2017 ต้นฉบับ "
            "จึงยังห้ามเคลมว่าสมการทั้งหมดเป็นไปตาม TGO"
        ),
        ANCHORS[35]: (
            f"[Implemented] {baseline_name} วิเคราะห์ local geometry เพื่อแยกไม้/ใบเป็นค่าเริ่มต้น "
            f"[Experimental] {candidate_name} เรียนรู้บริบทรอบจุดและมีผล Wan Wood IoU {wood_iou}, "
            f"Leaf IoU {leaf_iou}, Mean IoU {mean_iou}; ค่า Wood IoU ยังต่ำกว่า research target 0.70 "
            "และผลชุดนี้ไม่ใช่ independent final test"
        ),
        ANCHORS[36]: (
            "บทบาท: PyTorch และโมดูล ML ใช้กับ 3D point-cloud pipeline และ PointNet++ candidate "
            "การแปลงภาพถ่ายเป็น point cloud มีเพียง [Experimental] COLMAP/OpenMVS wrappers "
            "ที่ต้องพึ่ง external binaries; ยังไม่ใช่เส้นทางภาพถ่าย end-to-end ที่ตรวจสอบแล้ว"
        ),
        ANCHORS[37]: (
            "[Implemented in code/config] Platform: Supabase (PostgreSQL/Auth) "
            "โดยการใช้งานจริงขึ้นกับ environment และฐานข้อมูลที่ตั้งค่าไว้"
        ),
        ANCHORS[38]: (
            "Supabase client, JWT verification, schema/migrations และ RLS scripts มีใน repo "
            "แต่ reviewed core demo ใช้ local/shared filesystem สำหรับส่งไฟล์ระหว่าง API กับ worker "
            "จึงไม่ควรอ้างว่า Storage/DB ทุกเส้นทางผ่าน production deployment แล้ว"
        ),
        ANCHORS[39]: "[Planned production target] Platform: RunPod Serverless",
        ANCHORS[40]: (
            "RunPod handler และ GPU Dockerfile มีใน repo แต่ยังไม่มี continuously deployed production "
            "API/worker ที่ตรวจสอบแล้ว ปัจจุบัน demo ที่ยืนยันใช้ local worker ผ่าน temporary tunnel"
        ),
        ANCHORS[41]: (
            "[Implemented demo] ข้อมูล point cloud เข้าสู่ API และ local worker เพื่อรัน pipeline "
            "ขั้น 1-6 และ 8; ขั้น 7 Species Classification เป็น [Stub] "
            "[Planned] cloud GPU worker และการยืนยัน coefficients กับ TGO ยังไม่เสร็จ"
        ),
        ANCHORS[42]: (
            "[Implemented] การแสดงผลเชิงภาพมี 3D point-cloud viewer พร้อมสีแยก wood/leaf "
            "และสรุป DBH ความสูง carbon/CO2e [Planned] แผนที่พิกัดแปลงแบบ GIS ยังไม่พร้อมใช้"
        ),
        ANCHORS[43]: (
            "[Implemented in code] เว็บรองรับ Supabase email/password auth เมื่อมี environment และ API "
            "ตรวจ Bearer token สำหรับ jobs [Planned] GeoJSON plot ingestion, automatic hectare "
            "calculation และ reviewed Plot ID workflow ยังไม่ implement"
        ),
        ANCHORS[44]: (
            "[Implemented in code] Supabase session/JWT verification และ route protection มีบนเว็บ/API "
            "เมื่อกำหนด environment; RLS policies อยู่ใน SQL scripts แต่ต้อง deploy/verify กับฐานข้อมูลจริง "
            "ก่อนเคลมการบังคับใช้ครบทุกตารางและทุกบทบาท"
        ),
        ANCHORS[45]: (
            "[Partial schema] migration มีตารางผู้ใช้/แปลงและชนิดข้อมูลเชิงพื้นที่ "
            "แต่ plot ownership, GeoJSON ingestion และ spatial workflow ยังไม่อยู่ใน reviewed core path"
        ),
        ANCHORS[46]: (
            "[Partial schema] migration/model รองรับข้อมูลต้นไม้และค่าที่คำนวณได้บางส่วน "
            "แต่ tree endpoints ยังตอบ 501 และ workflow ตำแหน่ง/สถานะวางขายยังไม่ implement ครบ"
        ),
        ANCHORS[47]: (
            "[Implemented data file] species_db.csv เก็บ density และ allometric coefficients 5 รายการ "
            "[Planned verification] แหล่งที่มาและค่าทุกช่องยังต้องเทียบกับ TGO 2017/ต้นฉบับ "
            "ก่อนเคลมความแม่นยำตามชนิดพันธุ์"
        ),
        ANCHORS[48]: (
            "[Implemented] REST endpoints ที่ตรวจสอบแล้วครอบคลุม health, synchronous /upload/analyze, "
            "การดาวน์โหลด segmented cloud และรายชื่อชนิดพันธุ์; auth /me ตรวจ Supabase token "
            "[Stub/Planned] direct LAS/photo storage endpoints, tree retrieval, spatial filter, "
            "marketplace และ purchase endpoints ยังไม่พร้อม"
        ),
        ANCHORS[49]: (
            f"[Implemented geometry evidence] Demol isolated-tree 65 ต้นให้ DBH MAE {dbh_mae} cm, "
            f"Height MAE {height_mae} m และ Volume MAPE {volume_mape}% ภายใต้ 20,000-point cap, "
            "min-Z normalization และ tlsep ผลนี้ไม่ยืนยันว่าแม่นเท่าการวัดด้วยมือ "
            "และไม่ใช่ full-pipeline, biomass หรือ carbon validation"
        ),
        ANCHORS[50]: (
            f"[Experimental] synthetic-to-real zero-shot เป็นผลย้อนหลังโดยประมาณ (Mean IoU ~0.33) "
            f"ส่วน Wan same-environment best recorded คือ Wood IoU {wood_iou}, Leaf IoU {leaf_iou}, "
            f"Mean IoU {mean_iou} และ accuracy {accuracy}; held-out loader เดียวกันใช้เลือก best epoch "
            "งานถัดไปใช้ open datasets ตามคำแนะนำอาจารย์ ไม่ใช่เคลมว่าจะเก็บข้อมูลไม้ไทยเอง"
        ),
        ANCHORS[51]: (
            "[Implemented UI / Planned detail] รูปที่ 16 แสดงแนวทางหน้าจอสถานะ async job; "
            "ระบบจริงติดตามด้วย GET polling ส่วนความคืบหน้าแยก 8 ขั้น, WebSocket และ ETA "
            "ยังไม่ใช่ข้อมูล runtime ที่ worker ส่งครบ"
        ),
        ANCHORS[52]: (
            f"[Experimental] Wan real TLS best recorded มี Wood IoU {wood_iou}, Leaf IoU {leaf_iou}, "
            f"Mean IoU {mean_iou} และ accuracy {accuracy} บน held-out loader ที่ใช้เลือก best epoch ด้วย "
            "จึงยังไม่ใช่ independent final test และ PointNet++ ยังไม่ถูก promote"
        ),
        ANCHORS[53]: (
            "เป้าหมายที่ 13: prototype ช่วยสื่อสารและประมาณ carbon stock/CO2e จากต้นไม้ "
            "แต่ผลยังไม่ใช่ certified credit และยังไม่ผ่าน field/allometric validation ครบถ้วน "
            "จึงควรใช้เป็นเครื่องมือวิจัยและการเรียนรู้ ไม่ใช่หลักฐานชดเชยการปล่อย"
        ),
        ANCHORS[54]: (
            "เป้าหมายที่ 15: แนวคิด marketplace อาจสร้างแรงจูงใจให้ดูแลต้นไม้ในอนาคต "
            "แต่ marketplace, payment และ credit issuance เป็น [Planned] และยังไม่มีหลักฐานผลกระทบ "
            "ต่อการปลูกป่าหรือรายได้ของชุมชนจาก prototype ปัจจุบัน"
        ),
        ANCHORS[55]: (
            "ข้อจำกัดเรื่องพันธุ์ไม้และสมการคาร์บอน: species_db.csv มี 5 รายการ "
            "แต่ species classifier ยังเป็น [Stub] และ coefficients ยังไม่ผ่าน TGO verification "
            "โครงการยังไม่มี field benchmark ที่รองรับการระบุ error ของ carbon estimate เป็น 10-20%"
        ),
        ANCHORS[56]: (
            "ข้อพิจารณาเชิงชีววิทยาของไผ่: species_db.csv มีแถว Bambusa spp. และโค้ดคำนวณได้ "
            "แต่ความเหมาะสมของ culm-based equation, density และการใช้กับไผ่ชนิดต่าง ๆ "
            "ยังไม่ได้ verify กับเอกสารต้นฉบับ/TGO จึงต้องรายงานเป็นค่าที่ยังไม่รับรอง"
        ),
        ANCHORS[57]: (
            "ระยะเวลาในการประมวลผลขึ้นกับจำนวนจุด CPU/GPU และสภาพแวดล้อม "
            "ปัจจุบันยังไม่มี reproducible runtime benchmark ที่รองรับช่วง 5-30 นาทีต่อไฟล์ "
            "และ production worker deployment ยังเป็น [Planned]"
        ),
        ANCHORS[58]: (
            "[Implemented] AGB ใช้ species-specific equation a × DBH^b × H^c เมื่อพบชนิดไม้ "
            "และ coefficients ครบ; หากไม่พบจึงใช้ Chave 2014 pantropical fallback: "
            "AGB = 0.0673 × (ρ × DBH² × H)^0.976"
        ),
        ANCHORS[59]: (
            "[Implemented] คาร์บอนสะสม (kg C) = (AGB + BGB) × carbon_fraction โดย "
            "BGB = AGB × root_to_shoot; ค่า ratio/fraction มาจาก species_db.csv หรือค่า fallback "
            "จึงไม่ใช่ 1.24 × 0.47 คงที่สำหรับทุก species"
        ),
        ANCHORS[60]: (
            "หลักฐานข้อมูลในเล่มแยกเป็น 3 ขอบเขต: Demol 65 isolated trees สำหรับ geometry, "
            "synthetic held-out สำหรับ benchmark ภายใน และ Wan real TLS สำหรับ PointNet++ "
            "[Experimental] แต่ยังไม่มี independent final promotion split"
        ),
        ANCHORS[61]: (
            "6.5.2 [Experimental] ชุดข้อมูลและผล Wood-Leaf Segmentation (PointNet++)"
        ),
        ANCHORS[62]: (
            "รายละเอียดชุดข้อมูล: synthetic point clouds ใช้ฝึก/ทดสอบ benchmark ที่มี per-point labels "
            "ส่วนผลไม้จริงใช้ Wan et al. (2021); ทั้งสองขอบเขตไม่ใช่หลักฐาน production promotion"
        ),
    }


def _table_row_replacements(manifest: dict[str, Any]) -> dict[str, dict[int, str]]:
    validation = _require(manifest, "validation", "root")
    wan = _require(validation, "wan_held_out", "validation")
    synthetic = _require(validation, "synthetic_held_out", "validation")
    wood_iou = _require(wan, "wood_iou", "wan_held_out")
    leaf_iou = _require(wan, "leaf_iou", "wan_held_out")
    mean_iou = _require(wan, "mean_iou", "wan_held_out")
    synthetic_iou = _require(synthetic, "pointnet_mean_iou", "synthetic_held_out")

    return {
        TABLE_ROW_ANCHORS[0]: {
            0: "[Legacy approximate; provenance incomplete] PCA heuristic (zero-shot)",
            3: "~0.25",
        },
        TABLE_ROW_ANCHORS[1]: {
            0: "[Historical approximate] PointNet++ synthetic-to-real zero-shot",
            1: "~0.18",
            2: "~0.62",
            3: "~0.33",
        },
        TABLE_ROW_ANCHORS[2]: {
            0: "[Experimental] PointNet++ Wan held-out (loader also selected best epoch)",
            1: str(wood_iou),
            2: str(leaf_iou),
            3: str(mean_iou),
        },
        TABLE_ROW_ANCHORS[3]: {
            0: "[Experimental synthetic-only] PointNet++",
            1: "–",
            2: "–",
            3: str(synthetic_iou),
        },
        TABLE_ROW_ANCHORS[4]: {
            1: "[Implemented] Percentile-grid ground heuristic (CSF: reference only)"
        },
        TABLE_ROW_ANCHORS[5]: {
            1: "[Implemented] max-Z CHM + morphology (Pit-free: reference only)"
        },
        TABLE_ROW_ANCHORS[6]: {1: "[Implemented] Local maxima + Watershed"},
        TABLE_ROW_ANCHORS[7]: {
            1: "[Implemented default] tlsep; [Experimental] PointNet++"
        },
        TABLE_ROW_ANCHORS[8]: {
            1: "[Implemented] QSM-derived DBH/height/taper volume (not full TreeQSM)"
        },
        TABLE_ROW_ANCHORS[9]: {1: "[Stub] Species classifier; [Planned] ResNet/TFLite"},
        TABLE_ROW_ANCHORS[10]: {
            1: "[Implemented] species_db/Chave; [Planned verification] TGO 2017"
        },
        TABLE_ROW_ANCHORS[11]: {
            1: "[Implemented/configured] Supabase; [Planned production] RunPod"
        },
        TABLE_ROW_ANCHORS[12]: {1: "[Planned] Leaflet + GeoJSON GIS flow"},
        TABLE_ROW_ANCHORS[13]: {
            1: "[Experimental] Wan 2021 real TLS wood/leaf dataset"
        },
        TABLE_ROW_ANCHORS[14]: {
            1: "[Experimental wrappers] COLMAP + OpenMVS; external binaries required"
        },
    }


def _document_text(document: DocumentType) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(parts)


def _assert_truth_contract(
    document: DocumentType,
    manifest: dict[str, Any],
    table_replacements: dict[str, dict[int, str]],
) -> None:
    text = _document_text(document)
    forbidden_found = [claim for claim in FORBIDDEN_CLAIMS if claim in text]
    if forbidden_found:
        raise ValueError(
            "Forbidden legacy claim remains: "
            + "; ".join(repr(x) for x in forbidden_found)
        )

    validation = _require(manifest, "validation", "root")
    wan = _require(validation, "wan_held_out", "validation")
    demol = _require(validation, "demol_65", "validation")
    required = (
        "[Implemented]",
        "[Experimental]",
        "[Planned]",
        "[Stub]",
        f"Wood IoU {_require(wan, 'wood_iou', 'wan_held_out')}",
        f"Leaf IoU {_require(wan, 'leaf_iou', 'wan_held_out')}",
        f"Mean IoU {_require(wan, 'mean_iou', 'wan_held_out')}",
        f"accuracy {_require(wan, 'accuracy', 'wan_held_out')}",
        f"DBH MAE {_require(demol, 'dbh_mae_cm', 'demol_65')} cm",
        f"Height MAE {_require(demol, 'height_mae_m', 'demol_65')} m",
        f"Volume MAPE {_require(demol, 'volume_mape_pct', 'demol_65')}%",
        "not full TreeQSM",
        "certified carbon credit",
    )
    missing = [claim for claim in required if claim not in text]
    if missing:
        raise ValueError(
            "Required truth claim missing: " + "; ".join(repr(x) for x in missing)
        )

    rows = [row for table in document.tables for row in table.rows]
    for original_anchor, expected_cells in table_replacements.items():
        resulting_anchor = expected_cells.get(0, original_anchor)
        matches = [row for row in rows if row.cells[0].text == resulting_anchor]
        if len(matches) != 1:
            raise ValueError(
                f"Replaced table row must occur exactly once: {resulting_anchor!r}; "
                f"found {len(matches)}"
            )
        row = matches[0]
        for index, expected in expected_cells.items():
            if row.cells[index].text != expected:
                raise ValueError(
                    f"Table truth mismatch for {resulting_anchor!r}, cell {index}: "
                    f"expected {expected!r}, got {row.cells[index].text!r}"
                )


def _ensure_outside_repo(output: Path) -> None:
    repo_root = Path(__file__).resolve().parents[1]
    try:
        output.relative_to(repo_root)
    except ValueError:
        return
    raise ValueError("Report output must stay outside the repository")


def _assert_preserved(before: _Structure, after: _Structure) -> None:
    comparable = (
        ("paragraphs", before.paragraphs, after.paragraphs),
        ("tables", before.tables, after.tables),
        ("inline shapes", before.inline_shapes, after.inline_shapes),
        ("media files", before.media_files, after.media_files),
        ("sections", before.sections, after.sections),
    )
    changed = [f"{name}: {old} -> {new}" for name, old, new in comparable if old != new]
    if before.page_geometry != after.page_geometry:
        changed.append("page geometry changed")
    if changed:
        raise ValueError("DOCX structural preservation failed: " + "; ".join(changed))


def build_report(
    source: str | Path,
    output: str | Path,
    manifest: str | Path,
) -> ReportAudit:
    """Build a reviewed copy, aborting before output replacement on any ambiguity."""
    source_path = Path(source).resolve()
    output_path = Path(output).resolve()
    manifest_path = Path(manifest).resolve()
    if source_path == output_path:
        raise ValueError("Source and output must be different paths")
    _ensure_outside_repo(output_path)
    if not source_path.is_file():
        raise FileNotFoundError(source_path)
    if source_path.suffix.lower() != ".docx" or output_path.suffix.lower() != ".docx":
        raise ValueError("Source and output must both be .docx files")

    manifest_data = json.loads(manifest_path.read_text(encoding="utf-8"))
    replacements = _replacements(manifest_data)
    table_row_replacements = _table_row_replacements(manifest_data)
    if tuple(replacements) != ANCHORS:
        raise ValueError("Replacement map does not match the reviewed anchor contract")
    if tuple(table_row_replacements) != TABLE_ROW_ANCHORS:
        raise ValueError(
            "Table replacement map does not match the reviewed row contract"
        )

    source_sha256 = _sha256(source_path)
    source_document = Document(source_path)
    before = _structure(source_document, source_path)
    for anchor in ANCHORS:
        _validate_anchor(source_document, anchor)
    for anchor in TABLE_ROW_ANCHORS:
        _validate_table_row(source_document, anchor)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    file_descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{output_path.stem}-", suffix=".docx", dir=output_path.parent
    )
    os.close(file_descriptor)
    temporary_path = Path(temporary_name)
    try:
        shutil.copy2(source_path, temporary_path)
        document = Document(temporary_path)
        for anchor, replacement in replacements.items():
            replace_anchor(document, anchor, replacement)
        for anchor, row_replacements in table_row_replacements.items():
            row = _validate_table_row(document, anchor)
            _replace_table_row(row, row_replacements)
        document.save(temporary_path)

        reopened = Document(temporary_path)
        after = _structure(reopened, temporary_path)
        _assert_preserved(before, after)
        _assert_truth_contract(reopened, manifest_data, table_row_replacements)
        if _sha256(source_path) != source_sha256:
            raise ValueError("Source DOCX changed during report generation")

        output_sha256 = _sha256(temporary_path)
        os.replace(temporary_path, output_path)
        return ReportAudit(
            source_path=str(source_path),
            output_path=str(output_path),
            source_sha256=source_sha256,
            output_sha256=output_sha256,
            source_unchanged=_sha256(source_path) == source_sha256,
            anchors_replaced=len(replacements) + len(table_row_replacements),
            paragraph_anchors_replaced=len(replacements),
            table_row_anchors_replaced=len(table_row_replacements),
            truth_contract_passed=True,
            paragraphs_before=before.paragraphs,
            paragraphs_after=after.paragraphs,
            tables_before=before.tables,
            tables_after=after.tables,
            inline_shapes_before=before.inline_shapes,
            inline_shapes_after=after.inline_shapes,
            media_files_before=before.media_files,
            media_files_after=after.media_files,
            sections_before=before.sections,
            sections_after=after.sections,
            page_geometry_preserved=before.page_geometry == after.page_geometry,
        )
    finally:
        temporary_path.unlink(missing_ok=True)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--manifest", required=True, type=Path)
    arguments = parser.parse_args()
    audit = build_report(arguments.source, arguments.output, arguments.manifest)
    print(json.dumps(asdict(audit), ensure_ascii=True, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
