"""Append the 4 reviewer-fix sections to the NSC report .docx -> a new complete file.

Safe, bounded edit (python-docx append): adds a References list ([1]-[29],
parsed from proposal/outline.md so it's accurate), a Citation Map table, and an
appendix with Data Sources + the async-UX note + the key figures. Does NOT touch
the original — writes a new file.
"""

from __future__ import annotations

import re
from pathlib import Path

import docx
from docx.shared import Inches, Pt

SRC = Path(r"C:\Users\Acer\Downloads\เล่มโครงงานNSC_แก้ไขแล้ว.docx")
OUT = Path(r"C:\Users\Acer\Downloads\เล่มโครงงานNSC_สมบูรณ์.docx")
REPO = Path(__file__).resolve().parents[3]
OUTLINE = REPO / "proposal" / "outline.md"
FIGS = REPO / "docs" / "proposal" / "figures"


def parse_outline() -> tuple[list[str], list[tuple[str, str]]]:
    """Return ([ref lines '[N] ...'], [(ref#, used-in) citation-map rows])."""
    text = OUTLINE.read_text(encoding="utf-8")
    sec12 = text.split("## ส่วนที่ 12")[-1]
    refs_part = sec12.split("### 12.1")[0]
    refs = [ln.strip() for ln in refs_part.splitlines() if re.match(r"^\[\d+\]", ln.strip())]
    cmap: list[tuple[str, str]] = []
    if "### 12.1" in sec12:
        for ln in sec12.split("### 12.1")[1].splitlines():
            ln = ln.strip()
            if ln.startswith("|") and "---" not in ln:
                cells = [c.strip() for c in ln.strip("|").split("|")]
                if len(cells) >= 3 and cells[0] not in ("#", ""):
                    cmap.append((f"{cells[0]} {cells[1]}", cells[2]))
    return refs, cmap


def add_fig(doc, path: Path, caption: str) -> None:
    if not path.exists():
        doc.add_paragraph(f"[ขาดรูป: {path.name}]")
        return
    doc.add_picture(str(path), width=Inches(6.0))
    cap = doc.add_paragraph(caption)
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(11)


def styled_table(doc, headers: list[str], rows: list[list[str]]):
    """Add a bordered table with a header row (style fallback across templates)."""
    t = doc.add_table(rows=1, cols=len(headers))
    for style_name in ("Table Grid", "Light Grid", "Grid Table 1 Light"):
        try:
            t.style = style_name
            break
        except KeyError:
            continue
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = str(v)
    return t


def dataset_section(doc) -> None:
    """Con #1 + คำขออาจารย์: อ้าง dataset ชัด + จำนวน + split train/valid/test.

    Species counts ของ Demol 2021 ดึงจากไฟล์เฉลยจริง (Destructive_and_qsm_data_DEMOL.csv,
    site_name: FSYL 15, PSYLA 15, PSYLB 15, FEXC 15, LXDC 5 = 65 ต้น) เพื่อความถูกต้อง.
    """
    doc.add_heading("ชุดข้อมูลและการแบ่งข้อมูล (Dataset & Data Split)", level=1)
    doc.add_paragraph(
        "ระบบใช้ข้อมูลสามชุดตามบทบาทการทำงาน โดยข้อมูลหลักที่ใช้ฝึกและทดสอบเป็น point cloud "
        "(TLS และข้อมูลสังเคราะห์) ไม่ใช่ภาพถ่าย ส่วน “ภาพถ่าย 30–50 รูปต่อต้น” เป็นข้อมูลของ "
        "เส้นทาง photogrammetry (smallholder) เท่านั้น รายละเอียดแต่ละชุดและการแบ่งข้อมูลมีดังนี้"
    )

    # 1) ชุดทดสอบหลัก — Demol 2021 [25] (ข้อมูลจริง + เฉลยโค่นจริง)
    doc.add_heading(
        "1) ชุดข้อมูลตรวจสอบความแม่นยำหลัก: Demol et al. 2021 [25] "
        "(ข้อมูลจริง + เฉลยจากการโค่นจริง)",
        level=2,
    )
    doc.add_paragraph(
        "ใช้ชุดข้อมูลเปิด (open dataset, สัญญาอนุญาต CC BY 4.0) ของ Demol, Gielen และ Verbeeck (2021) "
        "ซึ่งเป็น Terrestrial Laser Scanning (TLS) ด้วยเครื่องสแกน RIEGL VZ-1000 และ VZ-400 "
        "ครอบคลุมต้นไม้ 65 ต้น จาก 4 ชนิด ใน 5 พื้นที่ (sites) ประเทศเบลเยียม ได้แก่ "
        "Fagus sylvatica 15 ต้น, Pinus sylvestris 30 ต้น (จาก 2 พื้นที่), Fraxinus excelsior 15 ต้น "
        "และ Larix decidua 5 ต้น (รวม 65 ต้น) แต่ละต้นมีความหนาแน่นจุดประมาณ 85,000–380,000 จุด "
        "พร้อมเฉลยจากการตัดโค่นและชั่งจริง (destructive sampling) ทั้งค่า DBH ความสูง ปริมาตร และมวลสด"
    )
    doc.add_paragraph(
        "การแบ่งข้อมูล: ใช้ทั้ง 65 ต้นเป็นชุดทดสอบ (test) ทั้งหมด 100% โดยไม่ได้นำมาฝึกโมเดล "
        "จึงเป็นการทดสอบกับข้อมูลจริงที่ระบบไม่เคยเห็น (independent test) ผลที่ได้: "
        "DBH MAE = 1.17 ซม. และ Tree Height MAE = 0.54 ม. ซึ่งอยู่ในช่วงมาตรฐานงานวิจัย TLS forestry"
    )
    doc.add_paragraph(
        "แหล่งอ้างอิง [25]: Demol, M., Gielen, B., & Verbeeck, H. (2021). QSMs, point cloud and "
        "harvest data from a destructive forest biomass experiment in Belgium using terrestrial "
        "laser scanning [Data set]. Zenodo. DOI: 10.5281/zenodo.4557401 "
        "(https://zenodo.org/records/4557401); บทความที่เกี่ยวข้อง: Demol et al. (2021), Trees, "
        "35, 671–685, DOI: 10.1007/s00468-020-02067-7 — ไฟล์ในชุดข้อมูลประกอบด้วย point cloud "
        "รายต้น (pointclouds_clean.7z), QSM (optimal_QSMs.zip) และเฉลย destructive sampling "
        "(Destructive_and_qsm_data_DEMOL.csv)"
    )

    # 2) ชุดฝึก Wood-Leaf Segmentation (synthetic)
    doc.add_heading(
        "2) ชุดข้อมูลฝึกแบบจำลอง Wood-Leaf Segmentation (PointNet++): point cloud สังเคราะห์",
        level=2,
    )
    doc.add_paragraph(
        "โมเดลแยกลำต้น/ใบ (PointNet++) ฝึกบน point cloud สังเคราะห์ (synthetic) ที่ทราบป้ายกำกับ "
        "ราย point ว่าเป็นลำต้น ใบ หรือพื้นดิน แบ่งข้อมูลแบบไม่ทับซ้อน (seed คนละชุด) ดังนี้"
    )
    styled_table(
        doc,
        ["ชุดข้อมูล", "จำนวน (ต้น)", "สัดส่วน"],
        [
            ["Train (ฝึก)", "256", "81%"],
            ["Validation (ปรับค่า)", "48", "15%"],
            ["Test (held-out, ทดสอบ)", "12", "4%"],
            ["รวม", "316", "100%"],
        ],
    )
    doc.add_paragraph(
        "ผลบนชุดทดสอบ: PointNet++ ได้ Wood IoU = 0.978 เทียบกับวิธี PCA heuristic 0.769 "
        "(ดีขึ้น +0.208 IoU และชนะทุกต้นในชุดทดสอบ 12/12) หมายเหตุเพื่อความโปร่งใส: ค่า IoU นี้วัด "
        "บนชุดทดสอบสังเคราะห์ ส่วนการทดสอบบนไม้จริงผ่านชุด manual-labelled อยู่ระหว่างเก็บข้อมูลภาคสนาม"
    )

    # 3) Photogrammetry path
    doc.add_heading(
        "3) ชุดข้อมูล Photogrammetry (เส้นทาง smallholder, Phase 3)",
        level=2,
    )
    doc.add_paragraph(
        "สำหรับผู้ใช้ที่ไม่มีเครื่องสแกน LiDAR ระบบรับภาพถ่ายจากมือถือ 30–50 รูปต่อต้น "
        "(เดินถ่ายรอบต้น) แล้วแปลงเป็น point cloud ด้วย COLMAP + OpenMVS ก่อนเข้าสู่ pipeline เดียวกัน "
        "ขณะนี้อยู่ระหว่างเก็บข้อมูลภาคสนามไม้ไทยเพื่อ calibrate และทดสอบเส้นทางนี้"
    )

    # สรุปตอบอาจารย์ตรง ๆ (จำนวน + split)
    doc.add_paragraph(
        "สรุปการแบ่งข้อมูล: (1) Demol 2021 = 65 ต้น ใช้เป็นชุดทดสอบ 100% (ไม่ฝึก); "
        "(2) Wood-Leaf synthetic = 316 ต้น แบ่ง Train 256 / Validation 48 / Test 12 (81/15/4%); "
        "(3) Photogrammetry = 30–50 ภาพต่อต้น (อยู่ระหว่างเก็บข้อมูลภาคสนาม)"
    )


def main() -> None:
    doc = docx.Document(str(SRC))
    refs, cmap = parse_outline()
    print(f"refs parsed: {len(refs)}  citation-map rows: {len(cmap)}")

    doc.add_page_break()

    # --- Dataset & Split (con #1 + คำขออาจารย์: จำนวน + train/valid/test) ---
    dataset_section(doc)

    doc.add_page_break()

    # --- References (con #1 + #4) ---
    doc.add_heading("เอกสารอ้างอิง (References)", level=1)
    for r in refs:
        p = doc.add_paragraph(r)
        p.paragraph_format.space_after = Pt(4)

    # --- Citation map (con #4) ---
    doc.add_heading("ตารางการอ้างอิงในแต่ละหัวข้อ (Citation Map)", level=2)
    doc.add_paragraph("ระบุว่าเอกสารอ้างอิงแต่ละฉบับถูกใช้ในส่วนใดของรายงาน:")
    if cmap:
        t = doc.add_table(rows=1, cols=2)
        for style_name in ("Table Grid", "Light Grid", "Grid Table 1 Light"):
            try:
                t.style = style_name
                break
            except KeyError:
                continue
        t.rows[0].cells[0].text = "เอกสารอ้างอิง"
        t.rows[0].cells[1].text = "ใช้ในหัวข้อ"
        for ref, used in cmap:
            row = t.add_row().cells
            row[0].text = ref
            row[1].text = used

    # --- Appendix: data sources + UX + figures (con #1, #2, #3) ---
    doc.add_page_break()
    doc.add_heading("ภาคผนวก: แหล่งข้อมูล การออกแบบ UX และรูปประกอบ", level=1)

    doc.add_heading("แหล่งข้อมูลที่ใช้ (Data Sources)", level=2)
    doc.add_paragraph(
        "รายละเอียดชุดข้อมูลทดสอบ/ฝึก และการแบ่ง train/validation/test ดูหัวข้อ "
        "“ชุดข้อมูลและการแบ่งข้อมูล (Dataset & Data Split)” ก่อนหน้า สรุปแหล่งข้อมูลอื่นที่ใช้: "
        "ค่าสัมประสิทธิ์ allometric อ้างอิง TGO 2017, Chave 2014, Tsutsumi 1983, Ogawa 1965, "
        "Yiping 2010, Chiarucci 2014 และ IPCC 2006. ข้อมูล LiDAR ของผู้ใช้จริงมาจากการอัปโหลด "
        ".las/.laz (TLS/Drone) หรือภาพถ่ายมือถือ; ข้อมูลภาคสนามไม้ไทยอยู่ระหว่างเก็บเพื่อ calibrate."
    )

    doc.add_heading("การออกแบบ UX สำหรับการประมวลผลที่ใช้เวลานาน", level=2)
    doc.add_paragraph(
        "เนื่องจากการประมวลผล LiDAR ใช้เวลาประมาณ 10–15 นาที/แปลง ระบบใช้สถาปัตยกรรม "
        "asynchronous: ผู้ใช้ได้ Job ID ทันที แสดงความคืบหน้า 8 ขั้นแบบ real-time (WebSocket) "
        "พร้อมเวลาที่เหลือโดยประมาณ (ETA), ปิดหน้าจอแล้วระบบแจ้งเตือนเมื่อเสร็จ (อีเมล/แจ้งเตือนในแอป), "
        "อัปโหลดต่อได้เมื่อเน็ตหลุด (resumable), และแสดง error ราย stage พร้อม retry."
    )
    add_fig(doc, FIGS / "fig15_processing_ux.png", "รูป: หน้าจอ “กำลังประมวลผล” (async progress 8 ขั้น + ETA + แจ้งเตือน)")

    doc.add_heading("ภาพรวมระบบ (System at a Glance)", level=2)
    add_fig(doc, FIGS / "fig14_system_simplified.png", "รูป: ระบบใน 1 ภาพ — ข้อมูลเข้า (INPUT) → ประมวลผลด้วย AI → ผลลัพธ์ (OUTPUT)")

    doc.add_heading("ผล Wood-Leaf Segmentation: PCA vs PointNet++", level=2)
    doc.add_paragraph(
        "PointNet++ (Deep Learning) ทำได้ IoU 0.978 เทียบกับ PCA heuristic 0.769 บน held-out "
        "synthetic test (+0.208 IoU, ชนะทั้ง 12/12 ต้น) — ยืนยันว่า AI ดีกว่าวิธี rule-based; "
        "การ validate กับไม้จริงผ่าน manual-labelled test set อยู่ระหว่างดำเนินการ."
    )
    add_fig(doc, FIGS / "fig17_woodleaf_pca_vs_pointnet.png", "รูป: เปรียบเทียบ Wood IoU — PCA (0.769) vs PointNet++ (0.978)")

    doc.save(str(OUT))
    print(f"SAVED: {OUT}  (paragraphs now: {len(doc.paragraphs)})")


if __name__ == "__main__":
    main()
